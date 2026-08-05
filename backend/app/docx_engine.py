import os
from docx import Document
from rapidfuzz import process, fuzz

KNOWN_LABELS = {
    "Employee Name": "employee_name",
    "Employee ID": "employee_id",
    "Department": "department",
    "Designation": "designation",
    "Month": "month",
    "Month & Year": "month_year",
    "Date of Joining": "date_of_joining",
    "Working Days": "working_days",
    "Paid Days": "paid_days",
    "Basic": "basic",
    "Basic Salary": "basic",
    "HRA": "hra",
    "Allowances": "allowances",
    "Bonus": "bonus",
    "Gross Salary": "gross",
    "EPF": "epf",
    "TDS": "tds",
    "LOP": "lop",
    "Net Pay": "net",
    "Net Salary": "net",
    "Bank": "bank",
    "Account Number": "account_number",
    "IFSC": "ifsc",
    "Payment Mode": "payment_mode",
    "Amount in Words": "amount_words"
}

def normalize_text(text):
    return text.strip().replace(":", "").replace("_", "").strip()

def extract_docx_fields(file_path):
    """
    Reads a DOCX file and uses fuzzy matching to find known labels.
    Returns a mapping JSON of where the editable values belong.
    """
    doc = Document(file_path)
    from collections import defaultdict
    mapping = defaultdict(list)
    
    label_keys = list(KNOWN_LABELS.keys())
    
    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = normalize_text(cell.text)
                if not text:
                    continue
                
                # Fuzzy match text against known labels
                match = process.extractOne(text, label_keys, scorer=fuzz.ratio, score_cutoff=85)
                if match:
                    matched_label, score, _ = match
                    field_id = KNOWN_LABELS[matched_label]
                    
                    # Assume the value is in the next column if available
                    if col_idx + 1 < len(row.cells):
                        mapping[field_id].append({
                            "type": "table_cell",
                            "table": table_idx,
                            "row": row_idx,
                            "column": col_idx + 1,
                            "label_found": text
                        })
    
    # Check paragraphs
    for p_idx, paragraph in enumerate(doc.paragraphs):
        text = normalize_text(paragraph.text)
        if not text:
            continue
            
        # e.g., "Employee Name : __________"
        for label_text, field_id in KNOWN_LABELS.items():
            norm_label = normalize_text(label_text)
            if text.lower().startswith(norm_label.lower()):
                mapping[field_id].append({
                    "type": "paragraph",
                    "paragraph": p_idx,
                    "label_found": label_text
                })

    return dict(mapping)


def fill_docx_template(template_path, mapping_json, values_dict, output_path):
    """
    Loads the template, injects values_dict based on mapping_json, and saves to output_path.
    """
    doc = Document(template_path)
    
    for field_id, location in mapping_json.items():
        if field_id not in values_dict:
            continue
            
        value = str(values_dict[field_id])
        
        # Support both old mapping format (dict) and new format (list of dicts)
        locations = location if isinstance(location, list) else [location]
        
        for loc in locations:
            if loc["type"] == "table_cell":
                table_idx = loc["table"]
                row_idx = loc["row"]
                col_idx = loc["column"]
                
                try:
                    cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
                    # Preserve formatting by replacing text in the first run and clearing the rest
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        if para.runs:
                            para.runs[0].text = value
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.add_run(value)
                except Exception as e:
                    print(f"Error filling table cell for {field_id}: {e}")
                    
            elif loc["type"] == "paragraph":
                p_idx = loc["paragraph"]
                try:
                    para = doc.paragraphs[p_idx]
                    label_text = loc.get("label_found", "")
                    
                    original_text = para.text
                    import re
                    # Match the label, optional spaces, optional colon/hyphen, optional spaces, and ANY number of underscores or dots
                    pattern = re.compile(re.escape(label_text) + r'\s*[:\-]?\s*[_\.]*', re.IGNORECASE)
                    
                    new_text = pattern.sub(label_text + ": " + value, original_text, count=1)
                    
                    if new_text != original_text:
                        # Keep first run's style to re-apply to the new single run if possible
                        style = para.runs[0].style if para.runs else None
                        para.text = new_text
                        if style and para.runs:
                            para.runs[0].style = style
                        
                except Exception as e:
                    print(f"Error filling paragraph for {field_id}: {e}")

    doc.save(output_path)
