import docx

doc = docx.Document()
section = doc.sections[0]
header = section.header
header.paragraphs[0].text = "TEST HEADER"

for sec in doc.sections:
    if sec.header:
        for p in sec.header.paragraphs:
            print("HEADER PARAGRAPH:", p.text)
