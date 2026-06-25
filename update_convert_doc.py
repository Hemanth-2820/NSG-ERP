import re

file_path = r"backend\app\routers\hr_portal.py"
with open(file_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Find the convert_doc block
pattern = r"import mammoth.*?@router\.post\(\"/onboarding/convert-doc\"\).*?return \{\"html\": html\}.*?raise HTTPException\(status_code=500, detail=str\(e\)\)"

new_code = """import mammoth
import docx

@router.post("/onboarding/convert-doc")
async def convert_doc(
    file: UploadFile = File(...),
    current_user: models.User = Depends(security.get_current_user)
):
    if current_user.role.lower() not in ["hr", "ceo", "admin", "superadmin"]:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
        
    try:
        content = await file.read()
        import io
        docx_file = io.BytesIO(content)
        
        # Extract main body
        result = mammoth.convert_to_html(docx_file)
        html_body = result.value
        
        # Try to extract headers/footers text
        try:
            docx_file.seek(0)
            doc = docx.Document(docx_file)
            header_html = ""
            footer_html = ""
            
            for section in doc.sections:
                if section.header:
                    header_text = "<br>".join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])
                    if header_text and header_text not in header_html:
                        header_html += f"<div style='border-bottom: 2px solid #ea580c; padding-bottom: 10px; margin-bottom: 20px; color: #0284c7; font-weight: bold; text-align: left; font-size: 18px;'>{header_text}</div>"
                
                if section.footer:
                    footer_text = "<br>".join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])
                    if footer_text and footer_text not in footer_html:
                        footer_html += f"<div style='border-top: 1px solid #d1d5db; padding-top: 10px; margin-top: 20px; font-size: 11px; text-align: center; color: #6b7280;'>{footer_text}</div>"
            
            final_html = f"{header_html}{html_body}{footer_html}"
        except Exception as ex:
            print("Error parsing header/footer:", ex)
            final_html = html_body
            
        return {"html": final_html}
    except Exception as e:
        print(f"Error converting docx: {e}")
        raise HTTPException(status_code=500, detail=str(e))"""

if re.search(pattern, content, re.DOTALL):
    content = re.sub(pattern, new_code, content, flags=re.DOTALL)
else:
    print("Could not find convert-doc block")
    
with open(file_path, 'w', encoding='utf-8') as f:
    f.write(content)

print("Updated convert-doc to include headers/footers")
